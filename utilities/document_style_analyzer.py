#!/usr/bin/env python3
"""
Document Style Analyzer
Allows users to upload example documents and configure header styles with filtering
"""

import sys
import os
from pathlib import Path
import tempfile
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field

# Add project root to Python path
project_root = Path(__file__).parent.parent  # utilities is in parent
sys.path.insert(0, str(project_root))

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
import fitz  # PyMuPDF

# Import the filtering utility
from .header_filter import HeaderFilter, ExactHeadingRule, ExactHeadingDetector
from utilities.header_filter import HeaderFilter  # ← FIXED: Added import

@dataclass
class TextStyle:
    """
    Represents a text style with formatting properties
    """
    font_size: Optional[float] = None
    is_bold: Optional[bool] = None
    is_italic: Optional[bool] = None
    font_name: Optional[str] = None
    text_color: Optional[str] = None
    style_name: Optional[str] = None  # DOCX style name like 'Heading 1'
    text_sample: str = ""
    # PDF-specific properties
    font_flags: Optional[int] = None  # PDF font flags
    font_family: Optional[str] = None  # PDF font family
    text_position: Optional[tuple] = None  # (x, y) coordinates
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "font_size": self.font_size,
            "is_bold": self.is_bold,
            "is_italic": self.is_italic,
            "font_name": self.font_name,
            "text_color": self.text_color,
            "style_name": self.style_name,
            "text_sample": self.text_sample,
            "font_flags": self.font_flags,
            "font_family": self.font_family,
            "text_position": self.text_position
        }

@dataclass
class HeaderAssignment:
    """
    Represents user assignment of a style to a header level with filtering
    """
    style: TextStyle
    level: int  # 1, 2, 3, etc.
    description: str = ""
    header_filter: HeaderFilter = field(default_factory=HeaderFilter)
    exact_heading_rules: List[ExactHeadingRule] = field(default_factory=list)

class DocumentStyleAnalyzer:
    """
    Analyzes document styles and allows user configuration
    """
    
    @staticmethod
    def extract_pdf_styles(pdf_path: str) -> List[TextStyle]:
        """
        Extract unique text styles directly from PDF (FAST method)
        """
        doc = fitz.open(pdf_path)
        styles = []
        seen_styles = set()
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_dict = page.get_text("dict")
            
            for block in text_dict["blocks"]:
                if "lines" in block:  # Text block
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if not text:
                                continue
                            
                            # Create style signature based on formatting properties
                            font_size = round(span["size"], 1)  # Round to avoid floating point issues
                            flags = span["flags"]
                            
                            # Determine bold/italic from flags
                            is_bold = bool(flags & 2**4)  # Bit 4 = bold
                            is_italic = bool(flags & 2**1)  # Bit 1 = italic
                            
                            style_signature = (
                                font_size,
                                is_bold,
                                is_italic,
                                span["font"],  # Font name
                                span["color"]  # Color
                            )
                            
                            if style_signature not in seen_styles:
                                style_info = TextStyle(
                                    font_size=font_size,
                                    is_bold=is_bold,
                                    is_italic=is_italic,
                                    font_name=span["font"],
                                    text_color=f"#{span['color']:06x}" if span['color'] != -1 else "black",
                                    font_flags=flags,
                                    font_family=span.get("font", "Unknown"),
                                    text_sample=text[:50],
                                    text_position=(span["bbox"][0], span["bbox"][1])  # x, y
                                )
                                styles.append(style_info)
                                seen_styles.add(style_signature)
        
        doc.close()
        return styles
    
    @staticmethod
    def extract_docx_styles(docx_path: str) -> List[TextStyle]:
        """
        Extract all unique text styles from DOCX document
        """
        doc = DocxDocument(docx_path)
        styles = []
        seen_styles = set()
        
        for para in doc.paragraphs:
            if para.text.strip():
                style_info = DocumentStyleAnalyzer._extract_paragraph_style(para)
                
                # Create a unique key for this style combination
                style_key = (
                    style_info.font_size,
                    style_info.is_bold,
                    style_info.is_italic,
                    style_info.font_name,
                    style_info.text_color,
                    style_info.style_name
                )
                
                if style_key not in seen_styles:
                    style_info.text_sample = para.text[:50]  # Sample text
                    styles.append(style_info)
                    seen_styles.add(style_key)
        
        return styles
    
    @staticmethod
    def _extract_paragraph_style(para: Paragraph) -> TextStyle:
        """
        Extract style information from a DOCX paragraph
        """
        style_info = TextStyle()
        
        # Get paragraph style name
        style_info.style_name = para.style.name if para.style else None
        
        # Get font properties from runs
        if para.runs:
            run = para.runs[0]  # Use first run for main properties
            font = run.font
            
            if font.size:
                style_info.font_size = font.size.pt if hasattr(font.size, 'pt') else font.size / 12700  # Convert from half-points
            
            style_info.is_bold = font.bold
            style_info.is_italic = font.italic
            style_info.font_name = font.name if font.name else "Unknown"
            
            # Extract color if available
            if font.color and font.color.rgb:
                style_info.text_color = f"#{font.color.rgb}"
        
        return style_info
    
    @staticmethod
    def analyze_document_styles(document_path: str) -> List[TextStyle]:
        """
        Analyze document styles - direct analysis for PDF, conversion for others
        """
        source_ext = Path(document_path).suffix.lower()
        
        if source_ext == '.pdf':
            # FAST: Direct PDF analysis without conversion
            return DocumentStyleAnalyzer.extract_pdf_styles(document_path)
        elif source_ext == '.docx':
            # Direct analysis
            return DocumentStyleAnalyzer.extract_docx_styles(document_path)
        elif source_ext == '.txt':
            # For TXT, create a temporary DOCX with basic patterns
            temp_docx_path = DocumentStyleAnalyzer._convert_txt_to_docx_basic(document_path)
            try:
                styles = DocumentStyleAnalyzer.extract_docx_styles(temp_docx_path)
                return styles
            finally:
                # Clean up temporary file
                if os.path.exists(temp_docx_path):
                    os.unlink(temp_docx_path)
        else:
            raise ValueError(f"Unsupported format: {source_ext}")
    
    @staticmethod
    def _convert_txt_to_docx_basic(txt_path: str) -> str:
        """
        Convert TXT to basic DOCX for style analysis (minimal processing)
        """
        from docx import Document
        import re
        
        doc = Document()
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split by paragraphs and apply basic pattern-based styling
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                lines = para_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        para = doc.add_paragraph(line)
                        
                        # Apply basic styling based on patterns
                        if re.match(r'^\d+\.(\d+\.?)*\s+.*$', line):
                            para.style = 'Heading 1' if hasattr(para.style, 'name') else None
                        elif re.match(r'^#+\s+.*$', line):
                            para.style = 'Heading 1' if hasattr(para.style, 'name') else None
                        elif line.upper() == line and len(line) > 5 and len(line) < 50:
                            para.style = 'Heading 1' if hasattr(para.style, 'name') else None
        
        # Create temporary DOCX file
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        
        doc.save(temp_docx.name)
        return temp_docx.name

def interactive_style_configuration(document_path: str) -> List[HeaderAssignment]:
    """
    Interactive function to let user configure header styles with filtering and exact headings
    """
    print(f"Analyzing document: {document_path}")
    print("Extracting text styles (fast analysis)...")
    
    styles = DocumentStyleAnalyzer.analyze_document_styles(document_path)
    
    print(f"\nFound {len(styles)} unique text styles:")
    print("="*60)
    
    for i, style in enumerate(styles, 1):
        print(f"{i:2d}. Style: {style.style_name or style.font_name or 'Unknown'}")
        print(f"    Font Size: {style.font_size}px" if style.font_size else "    Font Size: N/A")
        print(f"    Bold: {style.is_bold}" if style.is_bold is not None else "    Bold: N/A")
        print(f"    Italic: {style.is_italic}" if style.is_italic is not None else "    Italic: N/A")
        print(f"    Font: {style.font_name or style.font_family}" if style.font_name or style.font_family else "    Font: N/A")
        print(f"    Sample: {style.text_sample[:50]}...")
        print()
    
    assignments = []
    
    print("Configure header levels (Enter 0 to skip a style):")
    print("Header levels: 1 = Main, 2 = Section, 3 = Subsection, etc.")
    print("-" * 60)
    
    for i, style in enumerate(styles, 1):
        while True:
            try:
                level = input(f"Style {i} ('{style.style_name or style.font_name or style.font_family}') -> Header Level (1-5, 0 to skip): ").strip()
                level = int(level)
                
                if level == 0:
                    print("  → Skipped")
                    break
                elif 1 <= level <= 5:
                    description = input(f"  Description for Level {level} (optional): ").strip()
                    
                    # Ask for filtering options
                    filter_config = {
                        'include_words': [],
                        'exclude_words': [],
                        'include_regex': None,
                        'exclude_regex': None,
                        'min_length': None,
                        'max_length': None,
                        'starts_with': None,
                        'ends_with': None,
                        'contains_pattern': None
                    }
                    
                    # Include words filter
                    include_input = input(f"  Include words (comma-separated, optional): ").strip()
                    if include_input:
                        filter_config['include_words'] = [word.strip() for word in include_input.split(',') if word.strip()]
                    
                    # Exclude words filter
                    exclude_input = input(f"  Exclude words (comma-separated, optional): ").strip()
                    if exclude_input:
                        filter_config['exclude_words'] = [word.strip() for word in exclude_input.split(',') if word.strip()]
                    
                    # Include regex filter
                    include_regex_input = input(f"  Include regex pattern (optional): ").strip()
                    if include_regex_input:
                        filter_config['include_regex'] = include_regex_input
                    
                    # Exclude regex filter
                    exclude_regex_input = input(f"  Exclude regex pattern (optional): ").strip()
                    if exclude_regex_input:
                        filter_config['exclude_regex'] = exclude_regex_input
                    
                    # Length constraints
                    min_len_input = input(f"  Minimum length (optional): ").strip()
                    if min_len_input:
                        filter_config['min_length'] = int(min_len_input)
                    
                    max_len_input = input(f"  Maximum length (optional): ").strip()
                    if max_len_input:
                        filter_config['max_length'] = int(max_len_input)
                    
                    # Start/end patterns
                    starts_with_input = input(f"  Must start with (optional): ").strip()
                    if starts_with_input:
                        filter_config['starts_with'] = starts_with_input
                    
                    ends_with_input = input(f"  Must end with (optional): ").strip()
                    if ends_with_input:
                        filter_config['ends_with'] = ends_with_input
                    
                    header_filter = HeaderFilter(**filter_config)
                    
                    # Ask for exact heading rules
                    exact_rules = []
                    print(f"  Configure exact headings (leave empty to skip):")
                    while True:
                        exact_heading = input(f"    Exact heading text: ").strip()
                        if not exact_heading:
                            break
                        
                        try:
                            exact_level = int(input(f"    Level (1-5): ").strip() or str(level))
                            case_sensitive = input(f"    Case sensitive? (y/N): ").strip().lower() == 'y'
                            
                            rule = ExactHeadingRule(
                                heading_text=exact_heading,
                                level=exact_level,
                                case_sensitive=case_sensitive
                            )
                            exact_rules.append(rule)
                            print(f"      → Added exact rule for: '{exact_heading}' (Level {exact_level})")
                        except ValueError:
                            print("      → Invalid level, using original level")
                            rule = ExactHeadingRule(
                                heading_text=exact_heading,
                                level=level,
                                case_sensitive=False
                            )
                            exact_rules.append(rule)
                    
                    assignment = HeaderAssignment(
                        style=style,
                        level=level,
                        description=description or f"Level {level} Header",
                        header_filter=header_filter,
                        exact_heading_rules=exact_rules
                    )
                    assignments.append(assignment)
                    print(f"  → Assigned to Level {level}")
                    break
                else:
                    print("  → Please enter 1-5 or 0 to skip")
            except ValueError:
                print("  → Please enter a number")
    
    return assignments

def save_style_configuration(assignments: List[HeaderAssignment], config_path: str):
    """
    Save style configuration to JSON file with filtering options and exact headings
    """
    import json
    
    config_data = {
        "header_assignments": []
    }
    
    for assignment in assignments:
        assignment_data = {
            "level": assignment.level,
            "description": assignment.description,
            "style": assignment.style.to_dict(),
            "include_words": assignment.header_filter.include_words,
            "exclude_words": assignment.header_filter.exclude_words,
            "include_regex": assignment.header_filter.include_regex,
            "exclude_regex": assignment.header_filter.exclude_regex,
            "min_length": assignment.header_filter.min_length,
            "max_length": assignment.header_filter.max_length,
            "starts_with": assignment.header_filter.starts_with,
            "ends_with": assignment.header_filter.ends_with,
            "contains_pattern": assignment.header_filter.contains_pattern,
            "exact_heading_rules": [
                {
                    "heading_text": rule.heading_text,
                    "level": rule.level,
                    "case_sensitive": rule.case_sensitive,
                    "whole_word": rule.whole_word
                }
                for rule in assignment.exact_heading_rules
            ]
        }
        config_data["header_assignments"].append(assignment_data)
    
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config_data, f, indent=2, ensure_ascii=False)
    
    print(f"Configuration saved to: {config_path}")

def example_usage():
    """
    Example usage of the style analyzer with filtering and exact headings
    """
    print("Document Style Analyzer with Filtering and Exact Headings")
    print("="*50)
    
    # Get example document path
    doc_path = input("Enter path to example document (PDF/TXT/DOCX): ").strip()
    
    if not os.path.exists(doc_path):
        print(f"❌ File not found: {doc_path}")
        return
    
    # Analyze and configure styles with filtering and exact headings
    assignments = interactive_style_configuration(doc_path)
    
    if not assignments:
        print("No header styles configured. Exiting.")
        return
    
    # Save configuration
    config_path = input("Enter path to save configuration (e.g., header_config.json): ").strip()
    save_style_configuration(assignments, config_path)
    
    print(f"\n✅ Configuration complete!")
    print(f"Use this configuration file for batch processing.")
    
    print(f"\nConfiguration includes:")
    for assignment in assignments:
        print(f"  Level {assignment.level}: {assignment.description}")
        if assignment.header_filter.include_words or assignment.header_filter.exclude_words:
            print(f"    - Filtering: {len(assignment.header_filter.include_words)} include, {len(assignment.header_filter.exclude_words)} exclude")
        if assignment.exact_heading_rules:
            print(f"    - Exact headings: {len(assignment.exact_heading_rules)} rules")

if __name__ == "__main__":
    example_usage()