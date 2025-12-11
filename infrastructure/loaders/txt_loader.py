from typing import Dict, Any, Union, List, Optional
from domain.interfaces import IDocumentLoader
from domain.document import Document, Page, Section, DocumentFormat
from .docx.docx_loader import DocxLoader 
from infrastructure.processors.metadata_propagator import HeaderStyleDefinition, StyleBasedHeaderDetector
from utilities.header_filter import HeaderFilter
import os
import tempfile
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt

class TxtLoader(IDocumentLoader):
    """
    TXT document loader - converts to DOCX for style-based analysis
    """
    
    def __init__(self, header_style_definitions: Optional[List[HeaderStyleDefinition]] = None):
        self.header_style_definitions = header_style_definitions or []
        self.header_detector = StyleBasedHeaderDetector(self.header_style_definitions) if self.header_style_definitions else None
    
    def load(self, source: Union[str, Dict[str, Any]]) -> Document:
        """
        Load TXT by converting to DOCX and then using DOCX loader
        Args:
            source: File path string or configuration dict
        Returns:
            Document: Structured document (format preserved as TXT)
        """
        if isinstance(source, str):
            file_path = source
            config = {}
        else:
            file_path = source.get("path", "")
            config = source
        
        # Check for style definitions in config
        if "header_style_definitions" in config:
            self._update_style_definitions(config["header_style_definitions"])
        elif "style_config_path" in config:
            # Load from configuration file
            self._load_style_config(config["style_config_path"])
        
        # Convert TXT to temporary DOCX with potential header styling
        temp_docx_path = self._convert_txt_to_docx_with_patterns(file_path)
        
        try:
            # Use DOCX loader with same configuration
            docx_loader = DocxLoader(header_style_definitions=self.header_style_definitions)
            
            docx_source = {
                "path": temp_docx_path,
                "header_style_definitions": self.header_style_definitions
            }
            
            document = docx_loader.load(docx_source)
            # Update document format to original
            document.format = DocumentFormat.TXT
            
            return document
        finally:
            # Clean up temporary file
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
    
    def _convert_txt_to_docx_with_patterns(self, txt_path: str) -> str:
        """
        Convert TXT to DOCX while applying potential header patterns
        """
        
        doc = Document()
        
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split by paragraphs (multiple newlines)
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Split by single newlines to preserve line breaks
                lines = para_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # Create paragraph and potentially apply styling based on patterns
                        para = doc.add_paragraph()
                        run = para.add_run(line)
                        
                        # Apply potential header styling based on patterns
                        self._apply_potential_header_styling(run, line)
        
        # Create temporary DOCX file
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        
        doc.save(temp_docx.name)
        return temp_docx.name
    
    def _apply_potential_header_styling(self, run, text: str):
        """
        Apply potential header styling based on text patterns
        This helps with style detection after conversion
        """
        # Pattern for numbered headers (1., 1.1, 1.1.1, etc.)
        if re.match(r'^\d+\.(\d+\.?)*\s+.*$', text):
            # Likely a header - make it bold and larger
            run.font.bold = True
            run.font.size = Pt(14)  # Larger font for headers
        elif re.match(r'^#+\s+.*$', text):
            # Markdown-style headers
            run.font.bold = True
            run.font.size = Pt(16)
        elif re.match(r'^(?i)(chapter|section|part|appendix)\s+\d+.*$', text):
            # Academic-style headers
            run.font.bold = True
            run.font.size = Pt(15)
        elif text.upper() == text and len(text) > 5 and len(text) < 50:
            # ALL CAPS headers
            run.font.bold = True
            run.font.size = Pt(14)
        elif re.match(r'^[IVX]+\.?\s+.*$', text):
            # Roman numeral headers
            run.font.bold = True
            run.font.size = Pt(14)
    
    def supports_format(self, path: str) -> bool:
        """
        Check if loader supports TXT format
        """
        return path.lower().endswith('.txt')
    
    def get_document_metadata(self, path: str) -> Dict[str, Any]:
        """
        Extract TXT metadata without full loading
        """
        stat = os.stat(path)
        
        with open(path, 'r', encoding='utf-8') as f:
            line_count = sum(1 for line in f)
        
        return {
            "format": "TXT",
            "title": os.path.basename(path),
            "size": stat.st_size,
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "lines": line_count,
            "encoding": "utf-8"  # Assumed
        }
    
    def _get_file_size(self, path: str) -> int:
        """Get file size in bytes"""
        return os.path.getsize(path)
    
    def _update_style_definitions(self, style_configs: List[Dict[str, Any]]):
        """Update style definitions from configuration"""
        style_defs = []
        for config in style_configs:
            # Create header filter from config
            filter_config = {
                'include_words': config.get('include_words', []),
                'exclude_words': config.get('exclude_words', []),
                'include_regex': config.get('include_regex'),
                'exclude_regex': config.get('exclude_regex'),
                'min_length': config.get('min_length'),
                'max_length': config.get('max_length'),
                'starts_with': config.get('starts_with'),
                'ends_with': config.get('ends_with'),
                'contains_pattern': config.get('contains_pattern')
            }
            
            header_filter = HeaderFilter(**filter_config)
            
            style_def = HeaderStyleDefinition(
                level=config['level'],
                font_size=config.get('font_size'),
                is_bold=config.get('is_bold'),
                is_italic=config.get('is_italic'),
                starts_with_pattern=config.get('starts_with_pattern'),
                contains_pattern=config.get('contains_pattern'),
                header_filter=header_filter
            )
            style_defs.append(style_def)
        
        self.header_style_definitions = style_defs
        self.header_detector = StyleBasedHeaderDetector(self.header_style_definitions)
    
    def _load_style_config(self, config_path: str):
        """Load style configuration from JSON file"""
        import json
        
        with open(config_path, 'r', encoding='utf-8') as f:
            config_data = json.load(f)
        
        style_defs = []
        for item in config_data.get("header_assignments", []):
            style_data = item.get("style", {})
            
            # Create header filter from JSON config
            filter_config = {
                'include_words': item.get('include_words', []),
                'exclude_words': item.get('exclude_words', []),
                'include_regex': item.get('include_regex'),
                'exclude_regex': item.get('exclude_regex'),
                'min_length': item.get('min_length'),
                'max_length': item.get('max_length'),
                'starts_with': item.get('starts_with'),
                'ends_with': item.get('ends_with'),
                'contains_pattern': item.get('contains_pattern')
            }
            
            header_filter = HeaderFilter(**filter_config)
            
            style_def = HeaderStyleDefinition(
                level=item["level"],
                font_size=style_data.get("font_size"),
                is_bold=style_data.get("is_bold"),
                is_italic=style_data.get("is_italic"),
                starts_with_pattern=style_data.get("starts_with_pattern"),
                contains_pattern=style_data.get("contains_pattern"),
                header_filter=header_filter
            )
            style_defs.append(style_def)
        
        self.header_style_definitions = style_defs
        self.header_detector = StyleBasedHeaderDetector(self.header_style_definitions)