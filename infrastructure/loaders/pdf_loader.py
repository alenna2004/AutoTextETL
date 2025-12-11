from typing import Dict, Any, Union, List, Optional
from domain.interfaces import IDocumentLoader
from domain.document import Document, Page, Section, DocumentFormat
from .docx.docx_loader import DocxLoader
from infrastructure.processors.metadata_propagator import HeaderStyleDefinition, StyleBasedHeaderDetector
from utilities.header_filter import HeaderFilter
import fitz  # PyMuPDF
import os
import tempfile
from pathlib import Path

class PdfLoader(IDocumentLoader):
    """
    PDF document loader - converts to DOCX for style-based analysis
    """
    
    def __init__(self, header_style_definitions: Optional[List[HeaderStyleDefinition]] = None):
        self.header_style_definitions = header_style_definitions or []
        self.header_detector = StyleBasedHeaderDetector(self.header_style_definitions) if self.header_style_definitions else None
    
    def load(self, source: Union[str, Dict[str, Any]]) -> Document:
        """
        Load PDF by converting to DOCX and then using DOCX loader
        Args:
            source: File path string or configuration dict
        Returns:
            Document: Structured document (format preserved as PDF)
        """
        if isinstance(source, str):
            file_path = source
            config = {}
        else:
            file_path = source.get("path", "")
            config = source
        
        # Check for style definitions in config
        if "header_style_definitions" in config:
            self._update_style_definitions(config["header_style_definitions"])
        elif "style_config_path" in config:
            # Load from configuration file
            self._load_style_config(config["style_config_path"])
        
        # Convert PDF to temporary DOCX preserving styles
        temp_docx_path = self._convert_pdf_to_docx_with_styling(file_path)
        
        try:
            # Use DOCX loader with same configuration
            docx_loader = DocxLoader(header_style_definitions=self.header_style_definitions)
            
            docx_source = {
                "path": temp_docx_path,
                "header_style_definitions": self.header_style_definitions
            }
            
            document = docx_loader.load(docx_source)
            # Update document format to original
            document.format = DocumentFormat.PDF
            
            return document
        finally:
            # Clean up temporary file
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
    
    def _convert_pdf_to_docx_with_styling(self, pdf_path: str) -> str:
        """
        Convert PDF to DOCX while preserving text styles
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            # Get text with detailed formatting information
            text_dict = page.get_text("dict")
            
            for block in text_dict["blocks"]:
                if "lines" in block:  # Text block
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if text:
                                # Create paragraph with preserved styling
                                para = doc.add_paragraph()
                                
                                # Add text with preserved font properties
                                run = para.add_run(text)
                                
                                # Apply font properties
                                font_size = span["size"]
                                flags = span["flags"]
                                
                                # Set font size
                                if font_size:
                                    run.font.size = Pt(font_size)
                                
                                # Set bold (flag bit 2**4 = 16)
                                if flags & 16:
                                    run.font.bold = True
                                
                                # Set italic (flag bit 2**1 = 2)
                                if flags & 2:
                                    run.font.italic = True
                                
                                # Set alignment based on position
                                bbox = span["bbox"]
                                page_width = page.rect.width
                                text_width = bbox[2] - bbox[0]
                                left_margin = bbox[0]
                                
                                # If text is centered
                                if abs(left_margin - (page_width - text_width) / 2) < 10:
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # If text is right-aligned
                                elif abs(bbox[2] - page_width) < 10:
                                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                # If text is left-aligned
                                else:
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        pdf_doc.close()
        
        # Create temporary DOCX file
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        
        doc.save(temp_docx.name)
        return temp_docx.name
    
    def supports_format(self, path: str) -> bool:
        """
        Check if loader supports PDF format
        """
        return path.lower().endswith('.pdf')
    
    def get_document_metadata(self, path: str) -> Dict[str, Any]:
        """
        Extract PDF metadata without full loading
        """
        doc = fitz.open(path)
        metadata = doc.metadata
        page_count = len(doc)
        doc.close()
        
        return {
            "format": "PDF",
            "title": metadata.get("title", ""),
            "author": metadata.get("author", ""),
            "subject": metadata.get("subject", ""),
            "creator": metadata.get("creator", ""),
            "producer": metadata.get("producer", ""),
            "creation_date": metadata.get("creationDate", ""),
            "modification_date": metadata.get("modDate", ""),
            "page_count": page_count,
            "encrypted": metadata.get("encrypted", False),
            "file_size": self._get_file_size(path)
        }
    
    def _get_file_size(self, path: str) -> int:
        """Get file size in bytes"""
        return os.path.getsize(path)
    
    def _update_style_definitions(self, style_configs: List[Dict[str, Any]]):
        """Update style definitions from configuration"""
        style_defs = []
        for config in style_configs:
            # Create header filter from config
            filter_config = {
                'include_words': config.get('include_words', []),
                'exclude_words': config.get('exclude_words', []),
                'include_regex': config.get('include_regex'),
                'exclude_regex': config.get('exclude_regex'),
                'min_length': config.get('min_length'),
                'max_length': config.get('max_length'),
                'starts_with': config.get('starts_with'),
                'ends_with': config.get('ends_with'),
                'contains_pattern': config.get('contains_pattern')
            }
            
            header_filter = HeaderFilter(**filter_config)
            
            style_def = HeaderStyleDefinition(
                level=config['level'],
                font_size=config.get('font_size'),
                is_bold=config.get('is_bold'),
                is_italic=config.get('is_italic'),
                starts_with_pattern=config.get('starts_with_pattern'),
                contains_pattern=config.get('contains_pattern'),
                header_filter=header_filter
            )
            style_defs.append(style_def)
        
        self.header_style_definitions = style_defs
        self.header_detector = StyleBasedHeaderDetector(self.header_style_definitions)
    
    def _load_style_config(self, config_path: str):
        """Load style configuration from JSON file"""
        import json
        
        with open(config_path, 'r', encoding='utf-8') as f:
            config_data = json.load(f)
        
        style_defs = []
        for item in config_data.get("header_assignments", []):
            style_data = item.get("style", {})
            
            # Create header filter from JSON config
            filter_config = {
                'include_words': item.get('include_words', []),
                'exclude_words': item.get('exclude_words', []),
                'include_regex': item.get('include_regex'),
                'exclude_regex': item.get('exclude_regex'),
                'min_length': item.get('min_length'),
                'max_length': item.get('max_length'),
                'starts_with': item.get('starts_with'),
                'ends_with': item.get('ends_with'),
                'contains_pattern': item.get('contains_pattern')
            }
            
            header_filter = HeaderFilter(**filter_config)
            
            style_def = HeaderStyleDefinition(
                level=item["level"],
                font_size=style_data.get("font_size"),
                is_bold=style_data.get("is_bold"),
                is_italic=style_data.get("is_italic"),
                starts_with_pattern=style_data.get("starts_with_pattern"),
                contains_pattern=style_data.get("contains_pattern"),
                header_filter=header_filter
            )
            style_defs.append(style_def)
        
        self.header_style_definitions = style_defs
        self.header_detector = StyleBasedHeaderDetector(self.header_style_definitions)